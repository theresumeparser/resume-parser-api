from io import BytesIO

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from src.extraction.base import ExtractionError, ExtractionResult
from src.logging import get_logger

logger = get_logger(__name__)


def extract_docx(content: bytes, filename: str) -> ExtractionResult:
    try:
        doc = Document(BytesIO(content))
    except (PackageNotFoundError, Exception) as exc:
        logger.warning(
            "extraction_failed",
            method="docx",
            error=str(exc),
            filename=filename,
        )
        raise ExtractionError(
            f"Failed to open DOCX: {exc}", filename=filename
        ) from exc

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    paragraph_text = "\n".join(paragraphs)

    table_texts: list[str] = []
    for table in doc.tables:
        seen: set[str] = set()
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in seen:
                    seen.add(cell_text)
                    table_texts.append(cell_text)
    table_text = "\n".join(table_texts)

    parts = [p for p in (paragraph_text, table_text) if p]
    full_text = "\n\n".join(parts)

    result = ExtractionResult(
        text=full_text,
        pages=1,
        method="docx",
        source_filename=filename,
    )

    logger.info(
        "extraction_completed",
        method="docx",
        pages=result.pages,
        char_count=result.char_count,
        word_count=result.word_count,
        filename=filename,
    )

    return result
