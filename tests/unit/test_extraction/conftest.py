from io import BytesIO

import fitz
import pytest
from docx import Document


@pytest.fixture
def pdf_with_text() -> bytes:
    """Create a minimal 1-page PDF with selectable text."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "John Doe\nSoftware Engineer\nExperience and skills.")
    data = doc.tobytes()
    doc.close()
    return data


@pytest.fixture
def pdf_multi_page() -> bytes:
    """Create a 3-page PDF with text on each page."""
    doc = fitz.open()
    for i in range(3):
        page = doc.new_page()
        page.insert_text((72, 72), f"Page {i + 1} content for testing.")
    data = doc.tobytes()
    doc.close()
    return data


@pytest.fixture
def pdf_scanned() -> bytes:
    """Create a PDF with no selectable text (simulates a scanned doc)."""
    doc = fitz.open()
    doc.new_page()
    data = doc.tobytes()
    doc.close()
    return data


@pytest.fixture
def pdf_empty() -> bytes:
    """Minimal valid PDF with zero pages (hand-crafted)."""
    return (
        b"%PDF-1.0\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[]/Count 0>>endobj\n"
        b"xref\n0 3\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000052 00000 n \n"
        b"trailer<</Size 3/Root 1 0 R>>\n"
        b"startxref\n95\n%%EOF"
    )


@pytest.fixture
def docx_with_paragraphs() -> bytes:
    """Create a DOCX with body paragraphs."""
    doc = Document()
    doc.add_paragraph("John Doe")
    doc.add_paragraph("Software Engineer with 5 years of experience.")
    doc.add_paragraph("Skills: Python, FastAPI, Docker")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def docx_with_tables() -> bytes:
    """Create a DOCX with only tables (no paragraphs with content)."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "FastAPI"
    table.cell(1, 0).text = "Docker"
    table.cell(1, 1).text = "PostgreSQL"
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def docx_with_both() -> bytes:
    """Create a DOCX with both paragraphs and tables."""
    doc = Document()
    doc.add_paragraph("Jane Smith - Resume")
    doc.add_paragraph("Senior developer with extensive experience.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skill A"
    table.cell(0, 1).text = "Skill B"
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def docx_empty() -> bytes:
    """Create a DOCX with no content."""
    doc = Document()
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
